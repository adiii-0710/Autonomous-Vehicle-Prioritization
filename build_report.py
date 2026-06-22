from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Autonomous_Vehicle_Sensor_Data_Prioritization_Report.docx"
ASSETS = ROOT / "report_assets"
ASSETS.mkdir(exist_ok=True)

BLUE = "#1F4D78"
LIGHT_BLUE = "#E8EEF5"
GRAY = "#5B6573"
LIGHT_GRAY = "#F4F6F9"
RED = "#9B1C1C"
GREEN = "#245B45"


def font(size, bold=False):
    paths = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    if bold:
        paths.insert(0, "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def center_text(draw, box, text, use_font, fill):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=use_font, align="center", spacing=8)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((left + right - width) / 2, (top + bottom - height) / 2), text,
                        font=use_font, fill=fill, align="center", spacing=8)


def arrow(draw, start, end, fill=BLUE, width=7):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 18 * direction, y2 - 11), (x2 - 18 * direction, y2 + 11)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 11, y2 - 18 * direction), (x2 + 11, y2 - 18 * direction)]
    draw.polygon(points, fill=fill)


def create_architecture_diagram():
    path = ASSETS / "architecture_diagram.png"
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(44, True)
    label_font = font(28, True)
    small_font = font(22)
    center_text(draw, (0, 25, 1600, 95), "Autonomous Vehicle Sensor Data Prioritization Architecture", title_font, BLUE)

    sensors = [(85, 150, 330, 270, "LiDAR", "Distance"), (390, 150, 635, 270, "Radar", "Speed"),
               (695, 150, 940, 270, "Camera", "Object type"), (1000, 150, 1245, 270, "GPS", "Location")]
    for left, top, right, bottom, heading, detail in sensors:
        draw.rounded_rectangle((left, top, right, bottom), 18, fill=LIGHT_BLUE, outline=BLUE, width=4)
        center_text(draw, (left, top + 12, right, top + 72), heading, label_font, BLUE)
        center_text(draw, (left, top + 65, right, bottom - 5), detail, small_font, GRAY)

    fusion = (545, 380, 1055, 510)
    draw.rounded_rectangle(fusion, 20, fill="#DCEAF7", outline=BLUE, width=5)
    center_text(draw, (fusion[0], fusion[1] + 15, fusion[2], fusion[1] + 78), "Sensor Fusion Engine", label_font, BLUE)
    center_text(draw, (fusion[0], fusion[1] + 75, fusion[2], fusion[3] - 8), "Combines distance, speed, type, location, and direction", small_font, GRAY)

    risk = (545, 610, 1055, 740)
    draw.rounded_rectangle(risk, 20, fill="#FFF4E5", outline="#9A6700", width=5)
    center_text(draw, (risk[0], risk[1] + 15, risk[2], risk[1] + 78), "Risk Assessment Module", label_font, "#7A5A00")
    center_text(draw, (risk[0], risk[1] + 75, risk[2], risk[3] - 8), "Calculates risk score and collision probability", small_font, GRAY)

    quick = (1140, 390, 1510, 505)
    draw.rounded_rectangle(quick, 18, fill="#E9F4EC", outline=GREEN, width=5)
    center_text(draw, (quick[0], quick[1] + 15, quick[2], quick[1] + 67), "Quick Sort Engine", label_font, GREEN)
    center_text(draw, (quick[0], quick[1] + 63, quick[2], quick[3] - 6), "Ranks high risk first", small_font, GRAY)

    decision = (1140, 620, 1510, 735)
    draw.rounded_rectangle(decision, 18, fill="#FCEBEC", outline=RED, width=5)
    center_text(draw, (decision[0], decision[1] + 15, decision[2], decision[1] + 67), "Decision Control", label_font, RED)
    center_text(draw, (decision[0], decision[1] + 63, decision[2], decision[3] - 6), "Braking, speed reduction, monitoring", small_font, GRAY)

    for left, top, right, bottom, _, _ in sensors:
        arrow(draw, ((left + right) // 2, bottom), (800, 380))
    arrow(draw, (800, 510), (800, 610))
    arrow(draw, (1055, 675), (1140, 450))
    arrow(draw, (1325, 505), (1325, 620))
    image.save(path)
    return path


def create_priority_figure():
    path = ASSETS / "priority_ranking.png"
    image = Image.new("RGB", (1500, 840), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(40, True)
    header_font = font(27, True)
    body_font = font(27)
    center_text(draw, (0, 25, 1500, 85), "Quick Sort Priority Ranking - Sample Sensor Data", title_font, BLUE)
    x = [100, 320, 850, 1100]
    widths = [220, 530, 250, 260]
    headers = ["Priority", "Detected Object", "Risk Score", "Decision Level"]
    top = 150
    height = 92
    for index, heading in enumerate(headers):
        draw.rectangle((x[index], top, x[index] + widths[index], top + height), fill=BLUE)
        center_text(draw, (x[index], top, x[index] + widths[index], top + height), heading, header_font, "white")
    rows = [
        ("1", "Pedestrian", "88", "Emergency"),
        ("2", "Dog", "83", "High"),
        ("3", "Vehicle", "75", "High"),
        ("4", "Truck", "67", "Medium"),
        ("5", "Bicycle", "65", "Medium"),
    ]
    colors = ["#FDECEC", "#FFF3E4", "#FFF3E4", "#F4F6F9", "#F4F6F9"]
    for row_index, row in enumerate(rows):
        y = top + height * (row_index + 1)
        for col, value in enumerate(row):
            draw.rectangle((x[col], y, x[col] + widths[col], y + height), fill=colors[row_index], outline="#C7D1DC", width=2)
            center_text(draw, (x[col], y, x[col] + widths[col], y + height), value, body_font, "#1F2937")
    center_text(draw, (120, 705, 1380, 790), "The highest-risk object is sent to the Decision Control System first.", font(28, True), BLUE)
    image.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color="000000", italic=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_text(doc, text, style=None, align=None, after=6, before=0, bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        set_run(paragraph.add_run(bold_prefix), bold=True)
        set_run(paragraph.add_run(text[len(bold_prefix):]))
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = paragraph.add_run(text)
    set_run(run, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level == 1 else BLUE)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text))
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text))
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, heading in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(paragraph.add_run(heading), size=10, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 1 and len(headers) == 3 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(paragraph.add_run(str(value)), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(10)
    set_run(paragraph.add_run(text), size=9, italic=True, color=GRAY)


def add_code(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run(code)
    set_run(run, size=9.5, name="Courier New")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("AUTONOMOUS VEHICLE PRIORITIZATION SYSTEM"), size=8.5, bold=True, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Semester II - Data Structures and Algorithms | Page "), size=8.5, color=GRAY)
    add_page_number(footer)


def cover_page(doc, architecture_path):
    for _ in range(5):
        doc.add_paragraph()
    add_text(doc, "ITM SKILLS UNIVERSITY", align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("CASE STUDY REPORT"), size=22, bold=True, color=BLUE)
    add_text(doc, "on", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title.paragraph_format.space_after = Pt(18)
    set_run(main_title.add_run("AUTONOMOUS VEHICLE\nSENSOR DATA PRIORITIZATION SYSTEM"), size=24, bold=True, color="0B2545")
    doc.add_picture(str(architecture_path), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(doc, "Submitted by: Aditya Khare", align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_text(doc, "B.Tech CSE 2025-29 | Semester II", align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_text(doc, "Data Structures and Algorithms with C++", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def build_report():
    architecture_path = create_architecture_diagram()
    priority_path = create_priority_figure()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Autonomous Vehicle Sensor Data Prioritization System"
    doc.core_properties.author = "Aditya Khare"
    cover_page(doc, architecture_path)

    add_heading(doc, "1. Introduction")
    add_text(doc, "Autonomous vehicles must examine nearby objects and quickly decide which ones need immediate attention. "
                  "This project is a console-based C++ simulation of that prioritization process. It models sensor inputs, combines them through a Sensor Fusion Engine, evaluates risk, uses Quick Sort to rank objects, and produces a decision recommendation.")
    add_text(doc, "The report follows the case study architecture while keeping the implementation suitable for a Semester-II Data Structures and Algorithms project. It simulates sensor behavior; it does not connect to real vehicle hardware or control a real car.")

    add_heading(doc, "2. Objective of the Project")
    for item in [
        "Simulate LiDAR, Radar, Camera, and GPS readings for nearby road objects.",
        "Combine the readings using a Sensor Fusion Engine.",
        "Calculate a risk score automatically instead of accepting a manually entered score.",
        "Use manual Quick Sort to rank objects from highest risk to lowest risk.",
        "Display a Priority Ranking Module and a Decision Control recommendation.",
        "Validate distance, speed, and direction input values.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Technologies Used")
    add_table(doc, ["Area", "Details"], [
        ("Programming language", "C++17"),
        ("Core data structure", "vector<DetectedObject> and vector<SensorReading>"),
        ("Primary algorithm", "Manual Quick Sort with partition function"),
        ("Interface", "Menu-driven console application"),
        ("Documentation", "Microsoft Word report with architecture and ranking visuals"),
    ], [2600, 6760])

    add_heading(doc, "4. Autonomous Vehicle System Architecture")
    add_text(doc, "The project represents the complete architecture requested in the case study. Each sensor module supplies one part of the environmental picture, and the later modules convert that information into a prioritised vehicle response.")
    doc.add_picture(str(architecture_path), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 1: Simulated autonomous vehicle sensor-to-decision architecture.")
    add_table(doc, ["Module", "Project Role"], [
        ("LiDAR", "Supplies object distance in metres."),
        ("Radar", "Supplies relative speed in km/h."),
        ("Camera", "Identifies object type such as Pedestrian, Dog, Bicycle, Vehicle, or Truck."),
        ("GPS", "Stores a simulated road-location label."),
        ("Sensor Fusion Engine", "Combines readings into a DetectedObject."),
        ("Risk Assessment Module", "Calculates risk using distance, speed, type, direction, and collision probability."),
        ("Decision Control System", "Recommends braking, speed reduction, steering/lane evaluation, or monitoring."),
    ], [2850, 6510])

    add_heading(doc, "5. Data Structure Design")
    add_text(doc, "Two structures are used. SensorReading represents the raw simulated inputs. DetectedObject is the final fused record that the prioritization engine sorts.")
    add_code(doc, "struct SensorReading {\n    int id;\n    string cameraObjectName;\n    float lidarDistance;\n    float radarSpeed;\n    string gpsLocation;\n    string direction;\n};")
    add_code(doc, "struct DetectedObject {\n    int id;\n    string objectName;\n    float distance;\n    float speed;\n    int riskScore;\n};")
    add_text(doc, "The vector data structure stores multiple readings and provides index-based access required by the Quick Sort algorithm.")

    add_heading(doc, "6. Sensor Fusion and Risk Assessment")
    add_text(doc, "The Sensor Fusion Engine converts every SensorReading into a DetectedObject. The Risk Assessment Module then assigns a risk score from 0 to 100. It gives more importance to nearby objects, fast-moving objects, vulnerable object types, crossing or approaching movement, and higher predicted collision probability.")
    add_table(doc, ["Risk Input", "How It Affects Priority"], [
        ("LiDAR distance", "Shorter distance increases risk."),
        ("Radar speed", "Higher relative speed increases risk."),
        ("Camera object type", "Pedestrians, animals, bicycles, and motorcycles receive higher weights."),
        ("Movement direction", "Crossing and approaching objects receive additional risk."),
        ("Collision probability", "Estimated from distance, speed, and direction, then included in the final score."),
    ], [2850, 6510])

    add_heading(doc, "7. Quick Sort Prioritization Engine")
    add_text(doc, "Quick Sort is implemented manually through partition() and quickSort(). The last object in the current range is chosen as the pivot. Objects with a greater risk score are moved to the left of the pivot, which produces descending priority order.")
    add_code(doc, "if (objects[j].riskScore > pivot)\n{\n    i++;\n    swap(objects[i], objects[j]);\n}")
    add_text(doc, "The project does not use std::sort(). This keeps the focus on the Quick Sort data-structures-and-algorithms requirement.")

    add_heading(doc, "8. Priority Ranking and Decision Control")
    add_text(doc, "After Quick Sort, the Priority Ranking Module displays the highest-risk object first. The Decision Control System reads the first object and recommends a basic simulated vehicle action.")
    doc.add_picture(str(priority_path), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 2: Sample ranking produced after Sensor Fusion, Risk Assessment, and Quick Sort.")
    add_table(doc, ["Risk Score", "Decision Control Recommendation"], [
        ("85-100", "Apply emergency braking and activate collision warning."),
        ("65-84", "Reduce speed and evaluate steering or lane-change alternatives."),
        ("40-64", "Maintain safe distance and monitor the object closely."),
        ("0-39", "Continue normal driving while monitoring surroundings."),
    ], [2400, 6960])

    add_heading(doc, "9. Program Workflow")
    for item in [
        "Load sample sensor data or capture a new sensor reading.",
        "Read LiDAR distance, Radar speed, Camera object type, GPS label, and direction.",
        "Run the Sensor Fusion Engine to create detected objects.",
        "Run the Risk Assessment Module to calculate risk scores.",
        "Apply manual Quick Sort in descending risk-score order.",
        "Display the priority ranking and run the Decision Control System.",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "10. Time and Space Complexity Analysis")
    add_table(doc, ["Case", "Time Complexity"], [
        ("Best case", "O(n log n)"),
        ("Average case", "O(n log n)"),
        ("Worst case", "O(n^2)"),
    ], [3900, 5460])
    add_text(doc, "Quick Sort performs in-place partitioning. Its average recursive stack space is O(log n), so it uses low additional memory compared with algorithms that require a separate array.")

    add_heading(doc, "11. Results and Verification")
    add_text(doc, "The completed application was compiled and tested with the sample sensor scenario. The program displayed sensor-layer readings, fused object data, automatically calculated risk scores, sorted the objects manually using Quick Sort, and selected the highest-priority object for a decision recommendation.")
    add_table(doc, ["Verification Item", "Result"], [
        ("LiDAR, Radar, Camera, and GPS modules", "Simulated and displayed in the console."),
        ("Sensor Fusion Engine", "Creates DetectedObject records from sensor readings."),
        ("Risk Assessment Module", "Calculates risk automatically; manual risk entry is removed."),
        ("Quick Sort", "Manual implementation; no std::sort() used."),
        ("Priority Ranking", "Displays highest risk object first."),
        ("Decision Control", "Recommends an appropriate simulated vehicle action."),
    ], [4400, 4960])

    add_heading(doc, "12. Limitations")
    for item in [
        "Sensor values are simulated through sample data or user input; there is no physical sensor hardware.",
        "GPS is represented by a simple location label instead of real coordinates and maps.",
        "Collision probability is a simple educational formula, not a machine-learning prediction model.",
        "Decision outputs are console recommendations and do not control real vehicle actuators.",
        "The system is designed as an academic DSA demonstration, not a production autonomous-driving system.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. Conclusion")
    add_text(doc, "This project demonstrates how a complete autonomous-vehicle prioritization pipeline can be represented at a Semester-II level. Simulated sensor modules provide the input, Sensor Fusion and Risk Assessment convert readings into meaningful risk scores, and manual Quick Sort prioritizes the objects. The Decision Control System then focuses on the highest-risk object first. The project connects a classical sorting algorithm with a practical autonomous-driving use case while clearly remaining a simulation.")

    add_heading(doc, "References")
    add_text(doc, "1. Case Study 158: Autonomous Vehicle Sensor Data Prioritization System, ITM Skills University.")
    add_text(doc, "2. Project source code: Autonomous Vehicle Sensor Data Prioritization System (main.cpp).")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
